"""
CV Formatter — converts structured CV data into PDF or DOCX files.
Professional black and white layout — clean, recruiter-ready.
"""
import io
from datetime import datetime


def generate_cv_docx(cv_data: dict) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin   = Inches(0.85)
        section.right_margin  = Inches(0.85)

    personal  = cv_data.get("personal_info", {})
    full_name = personal.get("full_name", "")

    # ── Name ─────────────────────────────────────────────────────────────────
    # 16pt, bold, title case — not all-caps, not oversized
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run(full_name.title())
    name_run.bold = True
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

    # ── Contact row — split onto two lines if long ────────────────────────────
    contact_parts = []
    if personal.get("email"):    contact_parts.append(personal["email"])
    if personal.get("location"): contact_parts.append(personal["location"])
    if personal.get("phone"):    contact_parts.append(personal["phone"])

    link_parts = []
    if personal.get("linkedin"): link_parts.append(personal["linkedin"])
    if personal.get("github"):   link_parts.append(personal["github"])

    if contact_parts:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        cp.paragraph_format.space_after = Pt(1)
        cr = cp.add_run("  |  ".join(contact_parts))
        cr.font.size = Pt(9)
        cr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    if link_parts:
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lp.paragraph_format.space_after = Pt(4)
        lr = lp.add_run("  |  ".join(link_parts))
        lr.font.size = Pt(9)
        lr.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # ── Horizontal rule ───────────────────────────────────────────────────────
    rule_para = doc.add_paragraph()
    rule_para.paragraph_format.space_before = Pt(2)
    rule_para.paragraph_format.space_after  = Pt(6)
    pPr    = rule_para._p.get_or_add_pPr()
    pBdr   = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

    def add_section_header(title: str):
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(1)
        run = para.add_run(title.upper())
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        pPr    = para._p.get_or_add_pPr()
        pBdr   = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'),   'single')
        bottom.set(qn('w:sz'),    '4')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), '000000')
        pBdr.append(bottom)
        pPr.append(pBdr)
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(0)
        spacer.paragraph_format.space_after  = Pt(3)

    # ── Professional Summary ──────────────────────────────────────────────────
    summary = cv_data.get("professional_summary", "")
    if summary:
        add_section_header("Professional Summary")
        p = doc.add_paragraph(summary)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(10)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = cv_data.get("relevant_skills", {})
    if skills:
        add_section_header("Skills")
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if skill_list:
                    p   = doc.add_paragraph()
                    p.paragraph_format.space_after = Pt(2)
                    lbl = p.add_run(f"{category.title()}: ")
                    lbl.bold = True
                    lbl.font.size = Pt(10)
                    sl  = skill_list if isinstance(skill_list, list) else [skill_list]
                    val = p.add_run(", ".join(sl))
                    val.font.size = Pt(10)
        elif isinstance(skills, list):
            p = doc.add_paragraph(", ".join(skills))
            for r in p.runs:
                r.font.size = Pt(10)

    # ── Education ─────────────────────────────────────────────────────────────
    education = cv_data.get("education", [])
    if education:
        add_section_header("Education")
        for edu in education:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(f"{edu.get('degree', '')} in {edu.get('field', '')}")
            r.bold = True
            r.font.size = Pt(10)
            p2 = doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(3)
            r2 = p2.add_run(f"{edu.get('institution', '')}  —  {edu.get('year', '')}")
            r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            if edu.get("achievements"):
                p3 = doc.add_paragraph(edu["achievements"])
                p3.paragraph_format.space_after = Pt(3)
                for r3 in p3.runs:
                    r3.font.size = Pt(10)

    # ── Experience ────────────────────────────────────────────────────────────
    experience = cv_data.get("experience", [])
    if experience:
        add_section_header("Experience")
        for exp in experience:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(f"{exp.get('title', '')}  —  {exp.get('company', '')}")
            r.bold = True
            r.font.size = Pt(10)
            if exp.get("duration"):
                pd = doc.add_paragraph(exp["duration"])
                pd.paragraph_format.space_after = Pt(2)
                for rd in pd.runs:
                    rd.font.size = Pt(9)
                    rd.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
            for resp in exp.get("responsibilities", []):
                pr = doc.add_paragraph(style="List Bullet")
                pr.paragraph_format.space_after = Pt(1)
                rr = pr.add_run(resp)
                rr.font.size = Pt(10)
            for ach in exp.get("achievements", []):
                pa = doc.add_paragraph(style="List Bullet")
                pa.paragraph_format.space_after = Pt(1)
                ra = pa.add_run(f"✓ {ach}")
                ra.font.size = Pt(10)
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # ── Projects ──────────────────────────────────────────────────────────────
    projects = cv_data.get("projects", [])
    if projects:
        add_section_header("Projects")
        for proj in projects:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            r = p.add_run(proj.get("name", ""))
            r.bold = True
            r.font.size = Pt(10)
            if proj.get("description"):
                pd = doc.add_paragraph(proj["description"])
                pd.paragraph_format.space_after = Pt(1)
                for rd in pd.runs:
                    rd.font.size = Pt(10)
            if proj.get("technologies"):
                tech = proj["technologies"]
                tl   = tech if isinstance(tech, list) else [tech]
                pt   = doc.add_paragraph()
                pt.paragraph_format.space_after = Pt(2)
                rl   = pt.add_run("Tech Stack: ")
                rl.bold = True
                rl.font.size = Pt(10)
                rv   = pt.add_run(", ".join(tl))
                rv.font.size = Pt(10)
            if proj.get("link"):
                pl = doc.add_paragraph(proj["link"])
                pl.paragraph_format.space_after = Pt(3)
                for rl in pl.runs:
                    rl.font.size = Pt(9)
                    rl.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # ── Certifications ────────────────────────────────────────────────────────
    certifications = cv_data.get("certifications", [])
    if certifications:
        add_section_header("Certifications")
        for cert in certifications:
            pc = doc.add_paragraph(style="List Bullet")
            pc.paragraph_format.space_after = Pt(2)
            rc = pc.add_run(cert)
            rc.font.size = Pt(10)

    # ── Career Objective ──────────────────────────────────────────────────────
    objective = cv_data.get("career_objective", "")
    if objective:
        add_section_header("Career Objective")
        p = doc.add_paragraph(objective)
        p.paragraph_format.space_after = Pt(4)
        for r in p.runs:
            r.font.size = Pt(10)

    # ── Footer ────────────────────────────────────────────────────────────────
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.space_before = Pt(16)
    rf = pf.add_run(f"Generated by JobGad AI  —  {datetime.now().strftime('%B %Y')}")
    rf.font.size = Pt(8)
    rf.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def generate_cv_pdf(cv_data: dict) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor, black
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer,
        HRFlowable, ListFlowable, ListItem,
    )
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, pagesize=A4,
        rightMargin=20*mm, leftMargin=20*mm,
        topMargin=18*mm, bottomMargin=18*mm,
    )

    BLACK     = HexColor("#000000")
    DARKGRAY  = HexColor("#444444")
    LIGHTGRAY = HexColor("#999999")

    # ── Styles ────────────────────────────────────────────────────────────────
    # Name: 16pt, bold, title case — not oversized
    name_style = ParagraphStyle(
        "Name", fontSize=16, fontName="Helvetica-Bold",
        textColor=BLACK, alignment=TA_LEFT, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", fontSize=9, fontName="Helvetica",
        textColor=DARKGRAY, alignment=TA_LEFT, spaceAfter=2,
    )
    section_style = ParagraphStyle(
        "Section", fontSize=9, fontName="Helvetica-Bold",
        textColor=BLACK, spaceBefore=8, spaceAfter=2,
    )
    body_style = ParagraphStyle(
        "Body", fontSize=10, fontName="Helvetica",
        spaceAfter=3, leading=14,
    )
    bold_style = ParagraphStyle(
        "Bold", fontSize=10, fontName="Helvetica-Bold", spaceAfter=1,
    )
    small_gray = ParagraphStyle(
        "SmallGray", fontSize=9, fontName="Helvetica",
        textColor=DARKGRAY, spaceAfter=2,
    )
    footer_style = ParagraphStyle(
        "Footer", fontSize=8, fontName="Helvetica",
        textColor=LIGHTGRAY, alignment=TA_CENTER,
    )

    story    = []
    personal = cv_data.get("personal_info", {})

    # ── Name ──────────────────────────────────────────────────────────────────
    name = personal.get("full_name", "").title()
    story.append(Paragraph(name, name_style))

    # ── Contact row — split across two lines if needed ────────────────────────
    contact_parts = []
    if personal.get("email"):    contact_parts.append(personal["email"])
    if personal.get("location"): contact_parts.append(personal["location"])
    if personal.get("phone"):    contact_parts.append(personal["phone"])

    link_parts = []
    if personal.get("linkedin"): link_parts.append(personal["linkedin"])
    if personal.get("github"):   link_parts.append(personal["github"])

    if contact_parts:
        story.append(Paragraph("  |  ".join(contact_parts), contact_style))
    if link_parts:
        story.append(Paragraph("  |  ".join(link_parts), contact_style))

    # ── Divider ───────────────────────────────────────────────────────────────
    story.append(Spacer(1, 2))
    story.append(HRFlowable(width="100%", thickness=1.5, color=BLACK, spaceAfter=6))

    def add_section(title: str):
        story.append(Paragraph(title.upper(), section_style))
        story.append(HRFlowable(width="100%", thickness=0.5, color=BLACK, spaceAfter=4))

    # ── Professional Summary ──────────────────────────────────────────────────
    summary = cv_data.get("professional_summary", "")
    if summary:
        add_section("Professional Summary")
        story.append(Paragraph(summary, body_style))
        story.append(Spacer(1, 3))

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = cv_data.get("relevant_skills", {})
    if skills:
        add_section("Skills")
        if isinstance(skills, dict):
            for category, skill_list in skills.items():
                if skill_list:
                    sl = skill_list if isinstance(skill_list, list) else [skill_list]
                    story.append(Paragraph(
                        f"<b>{category.title()}:</b> {', '.join(sl)}", body_style
                    ))
        elif isinstance(skills, list):
            story.append(Paragraph(", ".join(skills), body_style))
        story.append(Spacer(1, 3))

    # ── Education ─────────────────────────────────────────────────────────────
    education = cv_data.get("education", [])
    if education:
        add_section("Education")
        for edu in education:
            story.append(Paragraph(
                f"{edu.get('degree','')} in {edu.get('field','')}", bold_style
            ))
            story.append(Paragraph(
                f"{edu.get('institution','')}  —  {edu.get('year','')}", small_gray
            ))
            if edu.get("achievements"):
                story.append(Paragraph(edu["achievements"], body_style))
            story.append(Spacer(1, 4))

    # ── Experience ────────────────────────────────────────────────────────────
    experience = cv_data.get("experience", [])
    if experience:
        add_section("Experience")
        for exp in experience:
            story.append(Paragraph(
                f"{exp.get('title','')}  —  {exp.get('company','')}", bold_style
            ))
            if exp.get("duration"):
                story.append(Paragraph(exp["duration"], small_gray))
            items = []
            for resp in exp.get("responsibilities", []):
                items.append(ListItem(Paragraph(resp, body_style), bulletColor=BLACK))
            for ach in exp.get("achievements", []):
                items.append(ListItem(Paragraph(f"✓ {ach}", body_style), bulletColor=BLACK))
            if items:
                story.append(ListFlowable(items, bulletType="bullet"))
            story.append(Spacer(1, 4))

    # ── Projects ──────────────────────────────────────────────────────────────
    projects = cv_data.get("projects", [])
    if projects:
        add_section("Projects")
        for proj in projects:
            story.append(Paragraph(proj.get("name", ""), bold_style))
            if proj.get("description"):
                story.append(Paragraph(proj["description"], body_style))
            if proj.get("technologies"):
                tech = proj["technologies"]
                tl   = tech if isinstance(tech, list) else [tech]
                story.append(Paragraph(f"<b>Tech Stack:</b> {', '.join(tl)}", body_style))
            if proj.get("link"):
                story.append(Paragraph(proj["link"], small_gray))
            story.append(Spacer(1, 4))

    # ── Certifications ────────────────────────────────────────────────────────
    certifications = cv_data.get("certifications", [])
    if certifications:
        add_section("Certifications")
        items = [
            ListItem(Paragraph(c, body_style), bulletColor=BLACK)
            for c in certifications
        ]
        story.append(ListFlowable(items, bulletType="bullet"))
        story.append(Spacer(1, 4))

    # ── Career Objective ──────────────────────────────────────────────────────
    objective = cv_data.get("career_objective", "")
    if objective:
        add_section("Career Objective")
        story.append(Paragraph(objective, body_style))
        story.append(Spacer(1, 4))

    # ── Footer ────────────────────────────────────────────────────────────────
    story.append(Spacer(1, 16))
    story.append(Paragraph(
        f"Generated by JobGad AI  —  {datetime.now().strftime('%B %Y')}",
        footer_style,
    ))

    doc.build(story)
    buffer.seek(0)
    return buffer.read()